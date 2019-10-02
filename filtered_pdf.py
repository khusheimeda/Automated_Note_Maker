import PyPDF2 
from docx import Document
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords

'''pdfFileObj = open('Samplings-Methods.pdf', 'rb') 
doc = Document()
para = doc.add_paragraph('')

pdfReader = PyPDF2.PdfFileReader(pdfFileObj) 


print(pdfReader.numPages) 
  

num_pages = pdfReader.numPages
text =''
for i in range(num_pages):
    pageObj = pdfReader.getPage(i)
    #print(pageObj.extractText()) 
    para.add_run(pageObj.extractText())
    text+=pageObj.extractText()

pdfFileObj.close() 
doc.save('C:/Users/drish/PES_UNI/Hackathon/pdf_to_word.docx')'''

doc = Document('C:/Users/drish/PES_UNI/Hackathon/pdf_to_word.docx')
bolds=[]
italics=[]
for para in doc.paragraphs:
    for run in para.runs:
        if run.italic :
            italics.append(run.text)
        if run.bold :
            bolds.append(run.text)
bolds = set(bolds)
bolds = list(bolds)
italics = set(italics)
italics = list(italics)

boltalic_Dict={'bold_phrases':bolds,
              'italic_phrases':italics}
print(boltalic_Dict)


