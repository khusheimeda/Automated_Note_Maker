from nltk.corpus import stopwords
from nltk.cluster.util import cosine_distance
import numpy as np
import networkx as nx
from textblob import TextBlob
import os
from nltk import word_tokenize, pos_tag
from docx.enum.text import WD_COLOR_INDEX
from docx import Document
import re
import sys

def read_article(file_name):
    file = open(file_name, "r")
    filedata = file.readlines()
    article = filedata[0].split(". ")
    sentences = []

    for sentence in article:
        # print(sentence)
        sentences.append(sentence.replace("[^a-zA-Z]", " ").split(" "))
    sentences.pop()

    return sentences


def sentence_similarity(sent1, sent2, stopwords=None):
    if stopwords is None:
        stopwords = []

    sent1 = [w.lower() for w in sent1]
    sent2 = [w.lower() for w in sent2]

    all_words = list(set(sent1 + sent2))

    vector1 = [0] * len(all_words)
    vector2 = [0] * len(all_words)

    # build the vector for the first sentence
    for w in sent1:
        if w in stopwords:
            continue
        vector1[all_words.index(w)] += 1

    # build the vector for the second sentence
    for w in sent2:
        if w in stopwords:
            continue
        vector2[all_words.index(w)] += 1

    return 1 - cosine_distance(vector1, vector2)


def build_similarity_matrix(sentences, stop_words):
    # Create an empty similarity matrix
    similarity_matrix = np.zeros((len(sentences), len(sentences)))

    for idx1 in range(len(sentences)):
        for idx2 in range(len(sentences)):
            if idx1 == idx2:  # ignore if both are same sentences
                continue
            similarity_matrix[idx1][idx2] = sentence_similarity(sentences[idx1], sentences[idx2], stop_words)

    return similarity_matrix


def generate_summary(file_name, top_n=5):
    stop_words = stopwords.words('english')
    summarize_text = []

    # Step 1 - Read text anc split it
    sentences = read_article(file_name)

    # Step 2 - Generate Similary Martix across sentences
    sentence_similarity_martix = build_similarity_matrix(sentences, stop_words)

    # Step 3 - Rank sentences in similarity martix
    sentence_similarity_graph = nx.from_numpy_array(sentence_similarity_martix)
    scores = nx.pagerank(sentence_similarity_graph)

    # Step 4 - Sort the rank and pick top sentences
    ranked_sentence = sorted(((scores[i], s) for i, s in enumerate(sentences)), reverse=True)
    # print("Indexes of top ranked_sentence order are ", ranked_sentence)

    for i in range(top_n):
        summarize_text.append(" ".join(ranked_sentence[i][1]))

    # Step 5 - Offcourse, output the summarize texr
    # print("Summarize Text: \n", ". ".join(summarize_text))
    return (". ".join(summarize_text))


# let's begin
concise = generate_summary("msft.txt", 2)
points = concise.split(",")
file_obj = open("filtered_msft.txt","w")
for i in points:
    file_obj.write(i)
    file_obj.write('\n')
file_obj.close()

f = open('filtered_msft.txt', 'r')
f1 = open('filtered_msft_1.txt', 'w')
k = 1
article = f.readlines()
for i in article:
    f1.write(str(k) + '.')
    f1.write(i + '\n')
    k+=1
f.close()
f1.close()
#os.remove('filtered_msft.txt')

document = Document()
myfile = open('filtered_msft_1.txt').read()
myfile = re.sub(r'[^\x00-\x7F]+|\x0c',' ', myfile) # remove all non-XML-compatible characters
p = document.add_paragraph(myfile)
document.save('C:/Users/KR/PycharmProjects/Automated_Note_Maker/filtered_msft_1'+'.docx')
#os.remove('filtered_msft_1.txt')

f = open('filtered_msft_1.txt')
l = f.read()
blob = TextBlob(l)
tagged_sent = pos_tag(l.split())
propernouns = [word for word,pos in tagged_sent if pos == 'NNP']

doc = Document('C:/Users/KR/PycharmProjects/Automated_Note_Maker/filtered_msft_1.docx')
#source = sys.argv[1]
l = l.split()
para = doc.add_paragraph('')
for i in l:
    if i not in propernouns:
        para.add_run(i)
    else:
        para.add_run(i).bold = True



"""for i in doc.paragraphs:
    txt = i.runs[0].text"""

doc.save('C:/Users/KR/PycharmProjects/Automated_Note_Maker/filtered_msft_1.docx')

'''for paragraph in document.paragraphs:
    for pn in propernouns:
        if pn in paragraph.text:
            for run in paragraph.runs:
                if pn in run.text:
                    x = run.text.split(pn)
                    run.clear()
                    for i in range(len(x)-1):
                        run.add_text(x[i])
                        run.add_text(pn)
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
'''

"""f = open('filtered_msft_1.txt')
l = f.read()
blob = TextBlob(l)
tagged_sent = pos_tag(l.split())
propernouns = [word for word,pos in tagged_sent if pos == 'NNP']
print(propernouns)
s = ''
for i in l:
    if i != ' ':
        s+=i
    else:
        if s in propernouns:
            #l.replace(s, "*" + s + '*')
            #l.replace(s, "\033[44;33mHello World!\033[m")
            l.replace(s, '\033[44;33m{}\033[m'.format(s))
        s = ''
f.close()
"""